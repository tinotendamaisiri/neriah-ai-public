"""Answer key management endpoints."""

from __future__ import annotations

import base64
import csv
import io
import logging
import time
import uuid
from datetime import datetime, timedelta, timezone

import google.auth
import google.auth.transport.requests
import requests as http
from flask import Blueprint, jsonify, request

from shared.auth import require_role
from shared.config import settings
from shared.firestore_client import delete_doc, get_doc, query, upsert
from shared.guardrails import log_ai_interaction, validate_input
from shared.gemma_client import (
    extract_answer_key_from_image,
    generate_marking_scheme,
    generate_marking_scheme_from_image,
    generate_scheme_from_text,
)
from shared.router import AIRequestType, route_ai_request
from shared.models import AnswerKey
from shared.observability import instrument_route
from shared.submission_codes import generate_unique_submission_code
from shared.user_context import get_user_context
from shared.weakness_tracker import update_student_weaknesses

logger = logging.getLogger(__name__)
answer_keys_bp = Blueprint("answer_keys", __name__)
homework_bp = Blueprint("homework", __name__)

# Maps MIME type → file extension for base64-encoded file inputs
_MEDIA_TYPE_EXT: dict[str, str] = {
    "image/jpeg": "jpg",
    "image/jpg": "jpg",
    "image/png": "png",
    "image/webp": "webp",
    "image/heic": "heic",
    "image/heif": "heif",
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword": "doc",
    "text/plain": "txt",
}


def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


# ─── Upload limits ────────────────────────────────────────────────────────────
_MAX_UPLOAD_BYTES = 10 * 1024 * 1024  # 10 MB

_ALLOWED_UPLOAD_EXTENSIONS: frozenset[str] = frozenset({
    "jpg", "jpeg", "png", "webp", "heic", "heif",
    "pdf", "docx", "doc", "txt",
})

_ALLOWED_MEDIA_TYPES: frozenset[str] = frozenset(_MEDIA_TYPE_EXT.keys())


def _validate_upload(file_bytes: bytes, filename: str) -> str | None:
    """
    Return an error string if the upload fails basic validation, else None.
    Checks file size and extension whitelist.
    """
    if len(file_bytes) > _MAX_UPLOAD_BYTES:
        return f"File too large (max {_MAX_UPLOAD_BYTES // (1024 * 1024)} MB)"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in _ALLOWED_UPLOAD_EXTENSIONS:
        return f"Unsupported file type '.{ext}'. Allowed: images, PDF, DOCX, TXT."
    return None


# ── Helpers ───────────────────────────────────────────────────────────────────

def _teacher_owns_class(teacher_id: str, class_id: str) -> bool:
    cls = get_doc("classes", class_id)
    return bool(cls and cls.get("teacher_id") == teacher_id)


def _normalise_question(raw: dict, idx: int) -> dict:
    """Map Gemma output (various field names) to AnswerKeyQuestion fields."""
    return {
        "question_number": int(raw.get("question_number") or raw.get("number") or idx + 1),
        "question_text": (raw.get("question_text") or raw.get("text") or "").strip(),
        "answer": (raw.get("answer") or raw.get("correct_answer") or "").strip(),
        "marks": float(raw.get("marks") or raw.get("max_marks") or 1),
        "marking_notes": raw.get("marking_notes"),
    }


def _extract_question_texts(qp_text: str) -> dict[int, str]:
    """
    Parse raw question paper text into a map of {question_number: question_text}.

    Handles formats like:
      1. What is ...        (dot after number)
      1) What is ...        (paren after number)
      Q1. What is ...       (Q prefix)
      Question 1: What is   (word prefix)

    Each question runs until the next numbered question or end of text.
    """
    import re
    # Match "1." or "1)" or "Q1." or "Q1)" or "Question 1:" etc. at start of line
    pattern = re.compile(
        r'^(?:Q(?:uestion)?\s*)?(\d{1,3})\s*[.):\-]\s*',
        re.MULTILINE | re.IGNORECASE,
    )
    matches = list(pattern.finditer(qp_text))
    if not matches:
        return {}

    result: dict[int, str] = {}
    for i, m in enumerate(matches):
        qnum = int(m.group(1))
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(qp_text)
        text = qp_text[start:end].strip()
        # Trim trailing blank lines
        text = re.sub(r'\n{2,}$', '', text)
        if text:
            result[qnum] = text
    return result


def _fill_empty_question_texts(questions: list[dict], qp_text: str | None) -> list[dict]:
    """
    Fill empty question_text fields by parsing the raw question paper text.
    Only fills questions where question_text is empty — never overwrites.
    Mutates and returns the same list.
    """
    if not qp_text:
        return questions
    has_empty = any(not q.get("question_text") for q in questions)
    if not has_empty:
        return questions

    extracted = _extract_question_texts(qp_text)
    if not extracted:
        return questions

    for q in questions:
        if not q.get("question_text"):
            qnum = q.get("question_number", 0)
            if qnum in extracted:
                q["question_text"] = extracted[qnum]
    return questions


def _extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """Extract plain text from PDF, DOCX, or TXT. Returns empty string on failure."""
    try:
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        if ext == "pdf":
            import pdfplumber
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                return "\n".join(page.extract_text() or "" for page in pdf.pages)
        elif ext in ("docx", "doc"):
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            return "\n".join(p.text for p in doc.paragraphs)
        elif ext == "txt":
            return file_bytes.decode("utf-8-sig", errors="replace").strip()
    except Exception:
        logger.exception("_extract_text_from_file failed for %s", filename)
    return ""


def _pdf_first_page_as_jpeg(pdf_bytes: bytes) -> bytes | None:
    """Render the first page of a PDF to JPEG bytes using pymupdf. Returns None on failure."""
    try:
        import fitz  # pymupdf
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        if doc.page_count == 0:
            return None
        page = doc[0]
        mat = fitz.Matrix(2.0, 2.0)  # 2× scale for better readability
        pix = page.get_pixmap(matrix=mat)
        return pix.tobytes("jpeg")
    except Exception:
        logger.warning("_pdf_first_page_as_jpeg failed", exc_info=True)
        return None


def _questions_from_file(
    file_bytes: bytes,
    filename: str,
    input_type: str = "question_paper",
    education_level: str = "",
    subject: str | None = None,
    user_ctx: dict | None = None,
) -> tuple[list[dict] | None, str | None, str | None]:
    """
    Process an uploaded file and return (questions, title_or_extracted_text, error).

    input_type:
      "question_paper" — teacher's question paper; Gemma generates answers (default)
      "answer_key"     — teacher's existing answer key; Gemma extracts Q&A

    Images → single Gemma multimodal call.
    PDF/DOCX/TXT → extract plain text → caller calls generate_marking_scheme.
    Returns (questions, title_or_text, error_str). questions is None for text-path files.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    image_exts = {"jpg", "jpeg", "png", "webp", "heic", "heif"}
    if ext in image_exts:
        if input_type == "question_paper":
            # Generate a full marking scheme from the question paper image
            scheme = generate_marking_scheme_from_image(
                file_bytes, education_level, subject, user_context=user_ctx,
            )
            if "error" in scheme:
                return None, None, scheme["error"]
            qs = scheme.get("questions", [])
            return [_normalise_question(q, i) for i, q in enumerate(qs)], scheme.get("title"), None
        else:
            # Extract Q&A from an existing answer key image
            result = extract_answer_key_from_image(file_bytes)
            qs = result.get("questions", [])
            return [_normalise_question(q, i) for i, q in enumerate(qs)], result.get("title"), None
    elif ext in ("pdf", "docx", "doc", "txt"):
        text = _extract_text_from_file(file_bytes, filename)
        if text and text.strip():
            return None, text, None  # caller calls generate_marking_scheme

        # Scanned PDF — no text layer. Convert first page to image and use Gemma vision.
        if ext == "pdf":
            logger.info("[answer_keys] PDF has no text layer — treating as scanned image")
            try:
                import fitz  # pymupdf
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                if doc.page_count > 0:
                    page = doc[0]
                    pix = page.get_pixmap(dpi=200)
                    img_bytes = pix.tobytes("jpeg")
                    doc.close()
                    logger.info("[answer_keys] Scanned PDF → JPEG image (%d bytes)", len(img_bytes))
                    scheme = generate_marking_scheme_from_image(
                        img_bytes, education_level, subject, user_context=user_ctx,
                    )
                    if "error" not in scheme:
                        qs = scheme.get("questions", [])
                        return [_normalise_question(q, i) for i, q in enumerate(qs)], scheme.get("title"), None
                    logger.warning("[answer_keys] Gemma failed on scanned PDF image: %s", scheme.get("error"))
                else:
                    doc.close()
            except Exception:
                logger.exception("[answer_keys] Scanned PDF image extraction failed")

            return None, None, "Could not read this PDF. Try taking a photo of the question paper instead."

        return None, text, None
    return None, None, f"Unsupported file type: .{ext}"


# ── Endpoints ─────────────────────────────────────────────────────────────────

@answer_keys_bp.get("/answer-keys")
@instrument_route("answer_keys.list", "answer_keys")
def list_answer_keys():
    teacher_id, err = require_role(request, "teacher")
    if err:
        return jsonify({"error": err}), 401

    class_id = request.args.get("class_id", "").strip()
    if not class_id:
        return jsonify({"error": "class_id query param is required"}), 400

    if not _teacher_owns_class(teacher_id, class_id):
        return jsonify({"error": "forbidden"}), 403

    results = query("answer_keys", [("class_id", "==", class_id)], order_by="created_at")

    # Auto-close any homework whose due_date has passed
    now_iso = _now_iso()
    for key in results:
        if (
            key.get("open_for_submission")
            and key.get("due_date")
            and key["due_date"] < now_iso
        ):
            upsert("answer_keys", key["id"], {"open_for_submission": False})
            key["open_for_submission"] = False

    # Enrich each homework with submission counts in one batch query.
    #
    # Three counts are exposed:
    #   - submission_count : every student submission, regardless of state.
    #   - graded_count     : AI has run grading (status in {graded, approved}
    #                        OR approved=True). Includes pending teacher review.
    #                        Kept for backwards compatibility.
    #   - approved_count   : teacher has explicitly approved (approved=True
    #                        OR status='approved'). Drives the homework
    #                        "Graded" badge — see HomeworkListScreen.getStatus.
    #   - pending_count    : neither AI- nor teacher- finished.
    all_subs = query("student_submissions", [("class_id", "==", class_id)])
    sub_counts: dict[str, dict] = {}
    EMPTY = {"submission_count": 0, "graded_count": 0, "approved_count": 0, "pending_count": 0}
    for sub in all_subs:
        ak_id = sub.get("answer_key_id", "")
        if not ak_id:
            continue
        if ak_id not in sub_counts:
            sub_counts[ak_id] = dict(EMPTY)
        sub_counts[ak_id]["submission_count"] += 1
        status = sub.get("status", "")
        is_approved = bool(sub.get("approved")) or status == "approved"
        is_ai_graded_or_approved = is_approved or status == "graded"
        if is_ai_graded_or_approved:
            sub_counts[ak_id]["graded_count"] += 1
        else:
            sub_counts[ak_id]["pending_count"] += 1
        if is_approved:
            sub_counts[ak_id]["approved_count"] += 1

    for key in results:
        counts = sub_counts.get(key["id"], dict(EMPTY))
        key.update(counts)

    return jsonify(results), 200


@answer_keys_bp.post("/answer-keys")
@instrument_route("answer_keys.create", "answer_keys")
def create_answer_key():
    teacher_id, err = require_role(request, "teacher")
    if err:
        return jsonify({"error": err}), 401

    is_multipart = "multipart" in (request.content_type or "")

    if is_multipart:
        class_id = (request.form.get("class_id") or "").strip()
        title = (request.form.get("title") or "").strip()
        education_level = (request.form.get("education_level") or "").strip()
        subject = (request.form.get("subject") or "").strip() or None
        question_paper_text = (request.form.get("question_paper_text") or "").strip()
        open_for_submission = request.form.get("open_for_submission", "false").lower() == "true"
        status = (request.form.get("status") or "").strip() or None
        input_type = (request.form.get("input_type") or "question_paper").strip()
        file = request.files.get("file")
        questions_raw = None
    else:
        body = request.get_json(silent=True) or {}
        class_id = (body.get("class_id") or "").strip()
        title = (body.get("title") or "").strip()
        education_level = (body.get("education_level") or "").strip()
        subject = (body.get("subject") or "").strip() or None
        question_paper_text = (body.get("question_paper_text") or "").strip()
        open_for_submission = bool(body.get("open_for_submission", False))
        status = body.get("status") or None
        input_type = (body.get("input_type") or "question_paper").strip()
        qs = body.get("questions")
        questions_raw = [_normalise_question(q, i) for i, q in enumerate(qs)] if qs else None
        file = None
        # Base64-encoded file (avoids multipart boundary issues on mobile)
        file_data_b64 = body.get("file_data")
        media_type = (body.get("media_type") or "").strip()

    if not class_id or not title:
        return jsonify({"error": "class_id and title are required"}), 400

    if not _teacher_owns_class(teacher_id, class_id):
        return jsonify({"error": "forbidden"}), 403

    # Auto-lookup education_level from class if not provided
    if not education_level:
        cls = get_doc("classes", class_id)
        education_level = (cls or {}).get("education_level", "")

    generated = False
    stored_qp_text: str | None = question_paper_text or None  # persisted for server-side regeneration

    # Build user context once — used by all three processing paths (file, base64, text).
    user_ctx = get_user_context(teacher_id, "teacher", class_id=class_id)

    qp_image_url: str | None = None  # question paper image URL in GCS

    # ── File upload processing ────────────────────────────────────────────────
    if file and file.filename:
        file_bytes = file.read()
        filename = (file.filename or "upload").lower()
        upload_err = _validate_upload(file_bytes, filename)
        if upload_err:
            return jsonify({"error": upload_err}), 400

        # Store the original question paper file to GCS
        try:
            import uuid as _uuid
            from shared.gcs_client import generate_signed_url, upload_bytes
            blob_name = f"question_papers/{teacher_id}/{_uuid.uuid4()}/{filename}"
            upload_bytes(settings.GCS_BUCKET_SCANS, blob_name, file_bytes, public=False)
            qp_image_url = generate_signed_url(settings.GCS_BUCKET_SCANS, blob_name, expiry_minutes=60 * 24 * 365)
            logger.info("[answer_keys] Stored question paper: %s", blob_name)
        except Exception:
            logger.warning("[answer_keys] Failed to store question paper to GCS — continuing without")

        qs_from_file, extracted_title_or_text, file_err = _questions_from_file(
            file_bytes, filename, input_type, education_level, subject, user_ctx,
        )

        if file_err:
            return jsonify({"error": file_err}), 400

        if qs_from_file is not None:
            # Image path — Gemma returned questions directly
            questions_raw = qs_from_file
            if not title or title == "Auto-generated scheme":
                title = extracted_title_or_text or title
            generated = True
        elif extracted_title_or_text:
            # Text path — need to call generate_marking_scheme
            question_paper_text = extracted_title_or_text
            stored_qp_text = question_paper_text

    elif not is_multipart and file_data_b64:
        # Base64 JSON path — mobile sends file as base64 string + media_type
        if media_type and media_type not in _ALLOWED_MEDIA_TYPES:
            return jsonify({"error": f"Unsupported media type '{media_type}'."}), 400
        try:
            file_bytes = base64.b64decode(file_data_b64)
        except Exception:
            return jsonify({"error": "Invalid base64 in file_data"}), 400

        # Store the original question paper to GCS
        try:
            import uuid as _uuid
            from shared.gcs_client import generate_signed_url, upload_bytes
            ext = _MEDIA_TYPE_EXT.get(media_type, "bin")
            blob_name = f"question_papers/{teacher_id}/{_uuid.uuid4()}/upload.{ext}"
            upload_bytes(settings.GCS_BUCKET_SCANS, blob_name, file_bytes, public=False)
            qp_image_url = generate_signed_url(settings.GCS_BUCKET_SCANS, blob_name, expiry_minutes=60 * 24 * 365)
            logger.info("[answer_keys] Stored question paper (b64): %s", blob_name)
        except Exception:
            logger.warning("[answer_keys] Failed to store question paper to GCS — continuing")
        if len(file_bytes) > _MAX_UPLOAD_BYTES:
            return jsonify({"error": f"File too large (max {_MAX_UPLOAD_BYTES // (1024 * 1024)} MB)"}), 400
        ext = _MEDIA_TYPE_EXT.get(media_type, "bin")
        filename = f"upload.{ext}"
        qs_from_file, extracted_title_or_text, file_err = _questions_from_file(
            file_bytes, filename, input_type, education_level, subject, user_ctx,
        )

        if file_err:
            return jsonify({"error": file_err}), 400

        if qs_from_file is not None:
            questions_raw = qs_from_file
            if not title or title == "Auto-generated scheme":
                title = extracted_title_or_text or title
            generated = True
        elif extracted_title_or_text:
            # Text path — need to call generate_marking_scheme
            question_paper_text = extracted_title_or_text
            stored_qp_text = question_paper_text

    # ── Text → generate marking scheme ───────────────────────────────────────
    if not questions_raw and question_paper_text:
        scheme = generate_marking_scheme(question_paper_text, education_level,
                                         user_context=user_ctx)
        questions_raw = [_normalise_question(q, i) for i, q in enumerate(scheme.get("questions", []))]
        if not title or title == "Auto-generated scheme":
            title = scheme.get("title") or title
        generated = True

    # Reject if no questions were generated — don't store incomplete answer keys.
    # The teacher must retry with a clearer image or paste the question text.
    if not questions_raw:
        logger.warning("[answer_keys] POST /answer-keys: no questions generated (title=%r)", title)
        return jsonify({
            "error": "Could not generate a marking scheme. "
                     "Try a clearer image or paste the question text instead.",
        }), 422

    # Fill empty question_text from question paper OCR — saves once, no Gemma call at read time.
    _fill_empty_question_texts(questions_raw, stored_qp_text)

    total_marks = sum(q.get("marks", 0) for q in questions_raw)
    # Reject grading-impossible homeworks. Sum-of-marks must be > 0 or the
    # AI grading pipeline has nothing to score against and the result is a
    # useless empty shell.
    if total_marks <= 0:
        return jsonify({
            "error": "Homework must have at least one question with marks greater than zero. "
                     "Add questions with their marks before saving.",
            "error_code": "INVALID_HOMEWORK_NO_MARKS",
        }), 400
    # Role invariant — refuse to credit homework creation to a student
    # account. teacher_id came from require_role(teacher) above so this
    # is belt-and-braces against a future bug that bypasses the JWT
    # check (e.g. a service-to-service caller passing a raw user_id).
    from shared.role_invariants import assert_is_teacher, RoleInvariantError
    try:
        assert_is_teacher(teacher_id)
    except RoleInvariantError as e:
        return jsonify({"error": str(e), "error_code": "ROLE_INVARIANT_VIOLATION"}), 400
    key = AnswerKey(
        class_id=class_id,
        teacher_id=teacher_id,
        title=title,
        education_level=education_level,
        subject=subject,
        questions=questions_raw,
        total_marks=total_marks,
        open_for_submission=open_for_submission,
        generated=generated,
        status=status,
        question_paper_text=stored_qp_text,
        submission_code=generate_unique_submission_code(),
    )
    doc = key.model_dump()
    if qp_image_url:
        doc["qp_image_url"] = qp_image_url
    upsert("answer_keys", key.id, doc)

    # Notify students if homework is open for submission
    if open_for_submission and class_id:
        try:
            from functions.push import notify_class_students
            notify_class_students(
                class_id,
                "New Homework",
                f"{title or subject or 'New assignment'} is ready — submit your work now.",
                {"screen": "StudentHome", "homework_id": key.id},
            )
        except Exception:
            logger.warning("Student notification failed (non-fatal)")

    return jsonify(doc), 201


@answer_keys_bp.put("/answer-keys/<key_id>")
@instrument_route("answer_keys.update", "answer_keys")
def update_answer_key(key_id: str):
    teacher_id, err = require_role(request, "teacher")
    if err:
        return jsonify({"error": err}), 401

    key = get_doc("answer_keys", key_id)
    if not key:
        return jsonify({"error": "Answer key not found"}), 404
    if key["teacher_id"] != teacher_id:
        return jsonify({"error": "forbidden"}), 403

    is_multipart = "multipart" in (request.content_type or "")
    updates: dict = {}

    if is_multipart:
        # File upload for auto-generating marking scheme on an existing key
        for field in ("title", "education_level", "subject", "due_date", "status"):
            val = (request.form.get(field) or "").strip()
            if val:
                updates[field] = val
        open_val = request.form.get("open_for_submission")
        if open_val is not None:
            updates["open_for_submission"] = open_val.lower() == "true"

        file = request.files.get("file")
        question_paper_text = (request.form.get("question_paper_text") or "").strip()

        # Build user context for AI calls in the update path
        upd_class_id = key.get("class_id") or ""
        upd_user_ctx = get_user_context(teacher_id, "teacher", class_id=upd_class_id)

        if file and file.filename:
            file_bytes = file.read()
            filename = (file.filename or "upload").lower()
            logger.info("[answer_keys] PUT file upload: %s (%d bytes)", filename, len(file_bytes))

            # Store the question paper to GCS
            try:
                import uuid as _uuid
                from shared.gcs_client import generate_signed_url, upload_bytes
                blob_name = f"question_papers/{teacher_id}/{_uuid.uuid4()}/{filename}"
                upload_bytes(settings.GCS_BUCKET_SCANS, blob_name, file_bytes, public=False)
                qp_url = generate_signed_url(settings.GCS_BUCKET_SCANS, blob_name, expiry_minutes=60 * 24 * 365)
                updates["qp_image_url"] = qp_url
            except Exception:
                logger.warning("[answer_keys] Failed to store question paper on PUT")

            education_level = updates.get("education_level") or key.get("education_level", "")
            qs_from_file, extracted_title_or_text, file_err = _questions_from_file(
                file_bytes, filename,
                education_level=education_level,
                user_ctx=upd_user_ctx,
            )

            if file_err:
                return jsonify({"error": file_err}), 400

            if qs_from_file is not None:
                updates["questions"] = qs_from_file
                updates["total_marks"] = sum(q.get("marks", 0) for q in qs_from_file)
                updates["generated"] = True
                updates["status"] = "draft"  # teacher must review before opening
            elif extracted_title_or_text:
                question_paper_text = extracted_title_or_text
                updates["question_paper_text"] = question_paper_text

        if not updates.get("questions") and question_paper_text:
            education_level = updates.get("education_level") or key.get("education_level", "")
            scheme = generate_marking_scheme(question_paper_text, education_level,
                                             user_context=upd_user_ctx)
            qs = [_normalise_question(q, i) for i, q in enumerate(scheme.get("questions", []))]
            if qs:
                updates["questions"] = qs
                updates["total_marks"] = sum(q.get("marks", 0) for q in qs)
                updates["generated"] = True
                updates["status"] = "draft"
                if not key.get("title") or key.get("title") == "Auto-generated scheme":
                    updates.setdefault("title", scheme.get("title") or key.get("title"))

    else:
        body = request.get_json(silent=True) or {}
        allowed_scalar = {"title", "education_level", "subject", "open_for_submission",
                          "due_date", "status", "generated", "total_marks"}
        updates = {k: v for k, v in body.items() if k in allowed_scalar}

        if "questions" in body:
            qs = body["questions"]
            updates["questions"] = [_normalise_question(q, i) for i, q in enumerate(qs)]
            updates["total_marks"] = sum(q.get("marks", 0) for q in updates["questions"])

        # Auto-generate from question_paper_text on a PUT
        question_paper_text = (body.get("question_paper_text") or "").strip()
        if not updates.get("questions") and question_paper_text:
            education_level = updates.get("education_level") or key.get("education_level", "")
            json_user_ctx = get_user_context(
                teacher_id, "teacher", class_id=key.get("class_id") or ""
            )
            scheme = generate_marking_scheme(question_paper_text, education_level,
                                             user_context=json_user_ctx)
            qs = [_normalise_question(q, i) for i, q in enumerate(scheme.get("questions", []))]
            if qs:  # Don't overwrite existing questions with empty Gemma output
                updates["questions"] = qs
                updates["total_marks"] = sum(q.get("marks", 0) for q in qs)
                updates["generated"] = True
                if not key.get("title") or key.get("title") == "Auto-generated scheme":
                    updates.setdefault("title", scheme.get("title") or key.get("title"))

    if not updates:
        logger.warning("[answer_keys] PUT /answer-keys/%s: no updates. is_multipart=%s has_file=%s",
                       key_id, is_multipart, bool(request.files.get("file") if is_multipart else False))
        return jsonify({"error": "Could not generate marking scheme from the uploaded file. Try a clearer image or paste the question text instead."}), 400

    # Fill empty question_text from question paper text before saving
    if "questions" in updates:
        qp = question_paper_text or key.get("question_paper_text") or ""
        _fill_empty_question_texts(updates["questions"], qp)

    # Block updates that would drive total_marks to zero. We only check when
    # total_marks is actually in this update (questions were touched) so
    # status/due-date-only updates on legacy zero-mark keys still go through.
    if "total_marks" in updates:
        try:
            new_total = float(updates["total_marks"])
        except (TypeError, ValueError):
            new_total = 0.0
        if new_total <= 0:
            return jsonify({
                "error": "Homework must have at least one question with marks greater than zero. "
                         "Add questions with their marks before saving.",
                "error_code": "INVALID_HOMEWORK_NO_MARKS",
            }), 400

    upsert("answer_keys", key_id, updates)

    # Notify students if homework was just opened for submission
    was_open = key.get("open_for_submission", False)
    now_open = updates.get("open_for_submission", was_open)
    if now_open and not was_open:
        cid = key.get("class_id", "")
        if cid:
            try:
                from functions.push import notify_class_students
                hw_title = updates.get("title") or key.get("title") or key.get("subject") or "Homework"
                notify_class_students(
                    cid,
                    "New Homework",
                    f"{hw_title} is now open — submit your work.",
                    {"screen": "StudentHome", "homework_id": key_id},
                )
            except Exception:
                logger.warning("Student notification on open failed (non-fatal)")

    return jsonify({**key, **updates}), 200


@answer_keys_bp.get("/answer-keys/<key_id>/questions")
@instrument_route("answer_keys.questions", "answer_keys")
def answer_key_questions(key_id: str):
    """
    Student-safe — return question texts and marks for an answer key.
    Does NOT return answers or marking_notes.
    Accessible by both teachers and students.
    """
    _, err = require_role(request, "teacher", "student")
    if err:
        return jsonify({"error": err}), 401

    key = get_doc("answer_keys", key_id)
    if not key:
        return jsonify({"error": "Answer key not found"}), 404

    questions = key.get("questions") or []
    qp_text = key.get("question_paper_text") or ""

    # Lazy backfill: if question_text is empty but question_paper_text exists,
    # parse and save so future reads are instant (no repeated parsing).
    has_empty = any(not q.get("question_text") for q in questions)
    if has_empty and qp_text:
        _fill_empty_question_texts(questions, qp_text)
        upsert("answer_keys", key_id, {"questions": questions})
        logger.info("[answer_keys] backfilled question_text for key=%s", key_id)

    out = [
        {
            "question_number": q.get("question_number", i + 1),
            "question_text": q.get("question_text", ""),
            "marks": q.get("marks", 0),
        }
        for i, q in enumerate(questions)
    ]
    result = {"questions": out}
    if qp_text:
        result["question_paper_text"] = qp_text
    return jsonify(result), 200


@answer_keys_bp.post("/answer-keys/<key_id>/close")
@instrument_route("answer_keys.close", "answer_keys")
def close_answer_key(key_id: str):
    """
    Close submissions for an answer key and trigger the Cloud Run batch grading job.
    POST /api/answer-keys/{id}/close
    """
    teacher_id, err = require_role(request, "teacher")
    if err:
        return jsonify({"error": err}), 401

    key = get_doc("answer_keys", key_id)
    if not key:
        return jsonify({"error": "Answer key not found"}), 404
    if key["teacher_id"] != teacher_id:
        return jsonify({"error": "forbidden"}), 403

    # Close submissions
    upsert("answer_keys", key_id, {"open_for_submission": False})

    # Count pending student submissions
    pending = query("student_submissions", [
        ("answer_key_id", "==", key_id),
        ("status", "==", "pending"),
    ])
    pending_count = len(pending)

    if pending_count > 0:
        _trigger_batch_grading_job(key_id)

    return jsonify({
        "message": "Submissions closed, grading started" if pending_count > 0 else "Submissions closed",
        "pending_count": pending_count,
    }), 200


def _trigger_batch_grading_job(answer_key_id: str) -> None:
    """Trigger the Cloud Run Job via the GCP REST API, injecting ANSWER_KEY_ID as an env override."""
    try:
        creds, project = google.auth.default(
            scopes=["https://www.googleapis.com/auth/cloud-platform"]
        )
        auth_req = google.auth.transport.requests.Request()
        creds.refresh(auth_req)

        job_name = (
            f"projects/{settings.GCP_PROJECT_ID}"
            f"/locations/{settings.GCP_REGION}"
            f"/jobs/{settings.CLOUD_RUN_JOB_NAME}"
        )
        url = f"https://run.googleapis.com/v2/{job_name}:run"

        payload = {
            "overrides": {
                "containerOverrides": [{
                    "env": [{"name": "ANSWER_KEY_ID", "value": answer_key_id}],
                }],
            },
        }

        resp = http.post(
            url,
            json=payload,
            headers={"Authorization": f"Bearer {creds.token}"},
            timeout=15,
        )
        resp.raise_for_status()
        logger.info("Batch grading job triggered for answer_key_id=%s", answer_key_id)
    except Exception:
        logger.exception("Failed to trigger batch grading job for answer_key_id=%s", answer_key_id)


@answer_keys_bp.delete("/answer-keys/<key_id>")
@instrument_route("answer_keys.delete", "answer_keys")
def delete_answer_key(key_id: str):
    teacher_id, err = require_role(request, "teacher")
    if err:
        return jsonify({"error": err}), 401

    key = get_doc("answer_keys", key_id)
    if not key:
        return jsonify({"error": "Answer key not found"}), 404
    if key["teacher_id"] != teacher_id:
        return jsonify({"error": "forbidden"}), 403

    delete_doc("answer_keys", key_id)
    return jsonify({"message": "deleted"}), 200


# ── Generate marking scheme from question paper image ─────────────────────────

@answer_keys_bp.post("/answer-keys/generate")
@instrument_route("answer_keys.generate", "answer_keys")
def generate_answer_key_scheme():
    """
    POST /api/answer-keys/generate

    Upload a question paper image and get a generated marking scheme back for
    review. Does NOT save to Firestore — the teacher reviews and edits the
    scheme, then saves it via POST /api/answer-keys.

    Request: multipart/form-data
      image          — question paper image file (required)
      education_level — e.g. "Form 2" (required)
      class_id        — the class this will be used for (required)
      subject         — optional subject name hint
    """
    teacher_id, err = require_role(request, "teacher")
    if err:
        return jsonify({"error": err}), 401

    image_file = request.files.get("image")
    if not image_file:
        return jsonify({"error": "image file is required"}), 400

    education_level = (request.form.get("education_level") or "").strip()
    class_id = (request.form.get("class_id") or "").strip()
    subject = (request.form.get("subject") or "").strip() or None

    if not education_level:
        return jsonify({"error": "education_level is required"}), 400
    if not class_id:
        return jsonify({"error": "class_id is required"}), 400

    if not _teacher_owns_class(teacher_id, class_id):
        return jsonify({"error": "forbidden"}), 403

    image_bytes = image_file.read()
    route_ai_request(AIRequestType.SCHEME)  # always AIRoute.CLOUD on the backend
    user_ctx = get_user_context(teacher_id, "teacher", class_id=class_id)
    scheme = generate_marking_scheme_from_image(image_bytes, education_level, subject,
                                                user_context=user_ctx)

    if "error" in scheme:
        return jsonify({"error": scheme["error"]}), 422

    return jsonify({"generated": True, "scheme": scheme}), 200


# ── POST /api/homework/generate-scheme ───────────────────────────────────────
#
# Short-form media_type values sent by the web demo ("image", "pdf", "word",
# "text") alongside full MIME types sent by the mobile app ("image/jpeg", etc.)

_MEDIA_TYPE_SHORT_TO_EXT: dict[str, str] = {
    "image": "jpg",
    "photo": "jpg",
    "pdf":   "pdf",
    "word":  "docx",
    "doc":   "docx",
    "text":  "txt",
}


@homework_bp.post("/homework/generate-scheme")
@instrument_route("homework.generate_scheme", "homework")
def create_homework_with_scheme():
    """
    POST /api/homework/generate-scheme

    Create a new homework (answer key) record AND generate its marking scheme
    via Gemma 4 in one call. Used by both the mobile app and the web demo.

    Body (JSON):
      title            str  — homework title (required)
      class_id         str  — owning class (required)
      subject          str  — subject name (optional)
      education_level  str  — e.g. "Form 2" (optional; looked up from class if absent)
      file_data        str  — base64-encoded file contents
      media_type       str  — short: "image" | "pdf" | "word" | "text"
                              OR full MIME: "image/jpeg", "application/pdf", etc.
      text             str  — raw question paper text (alias: question_paper_text)

    Returns the full answer_key document including questions[] and generated=true.
    """
    teacher_id, err = require_role(request, "teacher")
    if err:
        return jsonify({"error": err}), 401

    body = request.get_json(silent=True) or {}
    class_id             = (body.get("class_id") or "").strip()
    title                = (body.get("title") or "").strip()
    subject              = (body.get("subject") or "").strip() or None
    education_level      = (body.get("education_level") or "").strip()
    media_type           = (body.get("media_type") or "").strip()
    file_data_b64        = body.get("file_data")
    text_input           = (body.get("text") or body.get("question_paper_text") or "").strip()
    due_date             = (body.get("due_date") or "").strip() or None
    teacher_total_marks_raw = body.get("teacher_total_marks")
    max_total_marks: int | None = None
    if teacher_total_marks_raw is not None:
        try:
            max_total_marks = int(float(str(teacher_total_marks_raw)))
            if max_total_marks <= 0:
                max_total_marks = None
        except (ValueError, TypeError):
            max_total_marks = None

    if not class_id or not title:
        return jsonify({"error": "class_id and title are required"}), 400

    if not _teacher_owns_class(teacher_id, class_id):
        return jsonify({"error": "forbidden"}), 403

    # Auto-lookup education_level from class if not provided
    if not education_level:
        cls = get_doc("classes", class_id)
        education_level = (cls or {}).get("education_level", "Form 4")

    questions_raw: list[dict] | None = None
    stored_qp_text: str | None = None

    # ── File path (base64 encoded) ────────────────────────────────────────────
    if file_data_b64:
        # Accept both full MIME types (mobile) and short-form names (web demo)
        ext = (
            _MEDIA_TYPE_EXT.get(media_type)                      # "image/jpeg" → "jpg"
            or _MEDIA_TYPE_SHORT_TO_EXT.get(media_type.lower())  # "image" → "jpg"
            or "bin"
        )
        filename = f"upload.{ext}"
        logger.info(
            "create_homework_with_scheme: file_data present, media_type=%r ext=%r class_id=%s",
            media_type, ext, class_id,
        )
        try:
            file_bytes = base64.b64decode(file_data_b64)
        except Exception:
            return jsonify({"error": "Invalid base64 in file_data"}), 400

        user_ctx = get_user_context(teacher_id, "teacher", class_id=class_id)
        image_exts = {"jpg", "jpeg", "png", "webp", "heic", "heif"}

        if ext in image_exts:
            # ── Image: single multimodal Gemma call ───────────────────────────
            scheme = generate_marking_scheme_from_image(
                file_bytes, education_level, subject, user_context=user_ctx,
                max_total_marks=max_total_marks,
            )
            if "error" in scheme:
                return jsonify({"error": scheme["error"]}), 400
            raw_qs = scheme.get("questions", [])
            questions_raw = [_normalise_question(q, i) for i, q in enumerate(raw_qs)]
            if not title or title == "Auto-generated scheme":
                title = scheme.get("title") or title

        elif ext == "pdf":
            # ── PDF: try pdfplumber text extraction first ─────────────────────
            import pdfplumber
            pdf_text = ""
            try:
                with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                    pdf_text = "\n".join(page.extract_text() or "" for page in pdf.pages).strip()
                logger.info(
                    "create_homework_with_scheme: pdfplumber extracted %d chars from PDF",
                    len(pdf_text),
                )
            except Exception:
                logger.warning("pdfplumber failed for PDF upload", exc_info=True)

            if pdf_text:
                # Searchable PDF — use text path
                text_input = pdf_text
            else:
                # Scanned PDF — render first page to JPEG and use multimodal Gemma
                logger.info(
                    "create_homework_with_scheme: PDF text empty, falling back to image path",
                )
                jpeg_bytes = _pdf_first_page_as_jpeg(file_bytes)
                if jpeg_bytes:
                    scheme = generate_marking_scheme_from_image(
                        jpeg_bytes, education_level, subject, user_context=user_ctx,
                        max_total_marks=max_total_marks,
                    )
                    if "error" in scheme:
                        return jsonify({"error": scheme["error"]}), 400
                    raw_qs = scheme.get("questions", [])
                    questions_raw = [_normalise_question(q, i) for i, q in enumerate(raw_qs)]
                    if not title or title == "Auto-generated scheme":
                        title = scheme.get("title") or title
                else:
                    return jsonify({
                        "error": "Could not read this PDF. Try uploading a clearer scan or paste the question text instead.",
                    }), 422

        elif ext in ("docx", "doc"):
            # ── Word document: extract text ───────────────────────────────────
            from docx import Document
            try:
                doc_obj = Document(io.BytesIO(file_bytes))
                word_text = "\n".join(p.text for p in doc_obj.paragraphs).strip()
                logger.info(
                    "create_homework_with_scheme: Word doc extracted %d chars",
                    len(word_text),
                )
                if word_text:
                    text_input = word_text
            except Exception:
                logger.warning("Word doc extraction failed", exc_info=True)

        elif ext == "txt":
            text_input = file_bytes.decode("utf-8-sig", errors="replace").strip()

    # ── Text path ─────────────────────────────────────────────────────────────
    if not questions_raw and text_input:
        # Input guardrails on question paper text
        _valid_in, _cleaned_text = validate_input(text_input, role="teacher", max_tokens=4000)
        if not _valid_in:
            log_ai_interaction(
                teacher_id, "teacher", "generate-scheme", text_input, "",
                tokens_used=0, latency_ms=0, blocked=True, block_reason=_cleaned_text,
            )
            return jsonify({"error": _cleaned_text}), 403
        text_input = _cleaned_text

        stored_qp_text = text_input
        logger.info(
            "create_homework_with_scheme: text path, len=%d education_level=%r",
            len(text_input), education_level,
        )
        _t0 = time.time()
        qs_result, raw_response = generate_scheme_from_text(
            text_input, education_level, subject, max_total_marks=max_total_marks,
        )
        _latency_ms = int((time.time() - _t0) * 1000)
        if qs_result is None:
            log_ai_interaction(
                teacher_id, "teacher", "generate-scheme", text_input,
                raw_response or "", tokens_used=0, latency_ms=_latency_ms,
                blocked=False,
            )
            if raw_response:
                logger.error(
                    "create_homework_with_scheme: JSON parse failed. Raw: %s", raw_response,
                )
                return jsonify({
                    "error": "Gemma responded but the output could not be parsed as JSON.",
                    "raw": raw_response,
                }), 422
            return jsonify({
                "error": "Could not generate marking scheme. Try a clearer image or paste the question text instead.",
            }), 422
        log_ai_interaction(
            teacher_id, "teacher", "generate-scheme", text_input,
            str(qs_result), tokens_used=len(str(qs_result)) // 4,
            latency_ms=_latency_ms, blocked=False,
        )
        questions_raw = [_normalise_question(q, i) for i, q in enumerate(qs_result)]

    if not questions_raw:
        logger.warning(
            "create_homework_with_scheme: Gemma returned no questions "
            "(media_type=%r, has_file=%s, has_text=%s)",
            media_type, bool(file_data_b64), bool(text_input),
        )
        return jsonify({
            "error": "Could not generate marking scheme. "
                     "Try a clearer image or paste the question text instead.",
        }), 422

    # No hard cap on question count. The previous 10-question clamp
    # silently dropped questions 11+ on papers that had more, leaving
    # student marks scored out of the wrong total — every question Gemma
    # returns must be preserved. The Gemma prompt already biases toward
    # the natural question count of the source.

    # Fill empty question_text from OCR before saving
    _fill_empty_question_texts(questions_raw, stored_qp_text)

    # Default due_date: server sets now + 24 h if teacher didn't specify one
    if not due_date:
        due_date = (datetime.now(timezone.utc) + timedelta(hours=24)).isoformat()

    total_marks = max_total_marks or sum(q.get("marks", 0) for q in questions_raw)
    # Reject AI-generated shells with zero marks — Gemma can occasionally
    # return an empty/ malformed scheme, and we don't want to persist it.
    if total_marks <= 0:
        return jsonify({
            "error": "The generated scheme has no marks. Try uploading a clearer question paper, "
                     "or add questions and marks manually before saving.",
            "error_code": "INVALID_HOMEWORK_NO_MARKS",
        }), 400
    # Role invariant — same rationale as the manual-create path above.
    from shared.role_invariants import assert_is_teacher, RoleInvariantError
    try:
        assert_is_teacher(teacher_id)
    except RoleInvariantError as e:
        return jsonify({"error": str(e), "error_code": "ROLE_INVARIANT_VIOLATION"}), 400
    key = AnswerKey(
        class_id=class_id,
        teacher_id=teacher_id,
        title=title,
        education_level=education_level,
        subject=subject,
        questions=questions_raw,
        total_marks=float(total_marks),
        open_for_submission=False,
        generated=True,
        status="draft",
        question_paper_text=stored_qp_text,
        due_date=due_date,
        submission_code=generate_unique_submission_code(),
    )
    upsert("answer_keys", key.id, key.model_dump())
    logger.info(
        "create_homework_with_scheme: created key %s with %d questions (total_marks=%s due_date=%s)",
        key.id, len(questions_raw), total_marks, due_date,
    )
    return jsonify(key.model_dump()), 201


# ── Generate marking scheme from an already-saved homework image ─────────────

def _download_image_from_url(url: str) -> bytes:
    """Download image bytes from a gs:// or https:// URL."""
    if url.startswith("gs://"):
        from google.cloud import storage
        from shared.config import settings as _s
        parts = url[5:].split("/", 1)
        bucket_name, blob_path = parts[0], parts[1]
        client = storage.Client(project=_s.GCP_PROJECT_ID)
        return client.bucket(bucket_name).blob(blob_path).download_as_bytes()
    else:
        import google.auth
        import google.auth.transport.requests
        import requests as _http
        creds, _ = google.auth.default(scopes=["https://www.googleapis.com/auth/cloud-platform"])
        auth_req = google.auth.transport.requests.Request()
        creds.refresh(auth_req)
        resp = _http.get(url, headers={"Authorization": f"Bearer {creds.token}"}, timeout=30)
        resp.raise_for_status()
        return resp.content


@homework_bp.post("/homework/<homework_id>/generate-scheme")
@instrument_route("homework.regenerate_scheme", "homework")
def generate_scheme_from_homework(homework_id: str):
    """
    POST /api/homework/{homework_id}/generate-scheme

    Generate a marking scheme from the question paper image already stored on
    an existing homework (answer key) document. The homework must have a
    'question_paper_url' field pointing to a GCS image.

    Returns the generated scheme for review — does NOT save to Firestore.
    """
    teacher_id, err = require_role(request, "teacher")
    if err:
        return jsonify({"error": err}), 401

    homework = get_doc("answer_keys", homework_id)
    if not homework:
        return jsonify({"error": "Homework not found"}), 404
    if homework.get("teacher_id") != teacher_id:
        return jsonify({"error": "forbidden"}), 403

    question_paper_url = (homework.get("question_paper_url") or "").strip()
    if not question_paper_url:
        return jsonify({"error": "This homework has no stored question paper image. Upload the question paper to generate a scheme."}), 422

    education_level = homework.get("education_level") or ""
    if not education_level:
        cls = get_doc("classes", homework.get("class_id", ""))
        education_level = (cls or {}).get("education_level", "Form 4")

    subject = homework.get("subject") or None
    homework_class_id = homework.get("class_id") or ""

    try:
        image_bytes = _download_image_from_url(question_paper_url)
    except Exception:
        logger.exception("Failed to download question paper image for homework_id=%s", homework_id)
        return jsonify({"error": "Could not download the question paper image. Please try again."}), 502

    user_ctx = get_user_context(teacher_id, "teacher", class_id=homework_class_id)
    scheme = generate_marking_scheme_from_image(image_bytes, education_level, subject,
                                                user_context=user_ctx)

    if "error" in scheme:
        return jsonify({"error": scheme["error"]}), 422

    return jsonify({"generated": True, "scheme": scheme}), 200


# ── POST /api/answer-keys/<key_id>/regenerate ────────────────────────────────

@answer_keys_bp.post("/answer-keys/<key_id>/regenerate")
@instrument_route("answer_keys.regenerate", "answer_keys")
def regenerate_answer_key_scheme(key_id: str):
    """
    POST /api/answer-keys/{key_id}/regenerate

    Re-generate the marking scheme for a DRAFT answer key. Accepts the same
    inputs as create (multipart file or JSON text). Updates the draft in
    Firestore and returns the new questions so the teacher can review again.

    Falls back to the answer key's stored question_paper_text if no new input
    is provided (useful when the teacher just taps "Regenerate" without
    re-uploading).
    """
    teacher_id, err = require_role(request, "teacher")
    if err:
        return jsonify({"error": err}), 401

    key = get_doc("answer_keys", key_id)
    if not key:
        return jsonify({"error": "Answer key not found"}), 404
    if key.get("teacher_id") != teacher_id:
        return jsonify({"error": "forbidden"}), 403
    if key.get("status") != "draft":
        return jsonify({"error": "Only draft answer keys can be regenerated"}), 400

    education_level = key.get("education_level", "")
    subject = key.get("subject") or None
    class_id = key.get("class_id", "")
    questions_raw: list[dict] | None = None
    question_paper_text = ""

    is_multipart = "multipart" in (request.content_type or "")

    if is_multipart:
        education_level = (request.form.get("education_level") or education_level).strip()
        subject = (request.form.get("subject") or subject or "").strip() or None
        question_paper_text = (request.form.get("question_paper_text") or "").strip()
        file = request.files.get("file")

        if file and file.filename:
            file_bytes = file.read()
            filename = (file.filename or "upload").lower()
            user_ctx = get_user_context(teacher_id, "teacher", class_id=class_id)
            qs_from_file, extracted_or_text, file_err = _questions_from_file(
                file_bytes, filename, "question_paper", education_level, subject, user_ctx,
            )
            if file_err:
                return jsonify({"error": file_err}), 400
            if qs_from_file is not None:
                questions_raw = qs_from_file
            elif extracted_or_text:
                question_paper_text = extracted_or_text
    else:
        body = request.get_json(silent=True) or {}
        question_paper_text = (body.get("question_paper_text") or "").strip()
        education_level = (body.get("education_level") or education_level).strip()
        subject = (body.get("subject") or subject or "").strip() or None

        # Base64 file input
        file_data_b64 = body.get("file_data")
        media_type = (body.get("media_type") or "").strip()
        if file_data_b64 and not questions_raw:
            try:
                file_bytes = base64.b64decode(file_data_b64)
            except Exception:
                return jsonify({"error": "Invalid base64 in file_data"}), 400
            ext = _MEDIA_TYPE_EXT.get(media_type, "bin")
            filename = f"upload.{ext}"
            user_ctx = get_user_context(teacher_id, "teacher", class_id=class_id)
            qs_from_file, extracted_or_text, file_err = _questions_from_file(
                file_bytes, filename, "question_paper", education_level, subject, user_ctx,
            )
            if file_err:
                return jsonify({"error": file_err}), 400
            if qs_from_file is not None:
                questions_raw = qs_from_file
            elif extracted_or_text:
                question_paper_text = extracted_or_text

    # Fall back to stored question paper text
    if not questions_raw and not question_paper_text:
        question_paper_text = key.get("question_paper_text") or ""

    if not questions_raw and question_paper_text:
        scheme = generate_marking_scheme(question_paper_text, education_level)
        questions_raw = [_normalise_question(q, i) for i, q in enumerate(scheme.get("questions", []))]

    if not questions_raw:
        return jsonify({"error": "No question paper provided for regeneration and none stored on this answer key."}), 400

    total_marks = sum(q.get("marks", 0) for q in questions_raw)
    # Reject AI re-generation that comes back with zero marks — don't clobber
    # a working scheme with an empty one just because the LLM failed.
    if total_marks <= 0:
        return jsonify({
            "error": "Regeneration produced no marks. The existing scheme is unchanged. "
                     "Try a different question paper or edit the questions manually.",
            "error_code": "INVALID_HOMEWORK_NO_MARKS",
        }), 400
    upsert("answer_keys", key_id, {
        "questions": questions_raw,
        "total_marks": total_marks,
        "generated": True,
    })
    logger.info("regenerate_answer_key_scheme: regenerated %d question(s) for key %s", len(questions_raw), key_id)
    return jsonify({"questions": questions_raw, "total_marks": total_marks}), 200


# ── PATCH /api/homework/{id} ──────────────────────────────────────────────────

@homework_bp.patch("/homework/<homework_id>")
@instrument_route("homework.patch", "homework")
def patch_homework(homework_id: str):
    """
    Update homework fields.
    PATCH /api/homework/{homework_id}
    Body: { open_for_submission: bool, title: str, ... }
    """
    teacher_id, err = require_role(request, "teacher")
    if err:
        return jsonify({"error": err}), 401

    homework = get_doc("answer_keys", homework_id)
    if not homework:
        return jsonify({"error": "Homework not found"}), 404
    if homework.get("teacher_id") != teacher_id:
        return jsonify({"error": "forbidden"}), 403

    body = request.get_json(silent=True) or {}
    allowed = {"open_for_submission", "title", "subject", "education_level", "due_date", "status"}
    updates = {k: v for k, v in body.items() if k in allowed}

    if not updates:
        return jsonify({"error": "No updatable fields provided"}), 400

    upsert("answer_keys", homework_id, updates)
    return jsonify({**homework, **updates}), 200


# ── POST /api/homework/{id}/grade-all ─────────────────────────────────────────

@homework_bp.post("/homework/<homework_id>/grade-all")
@instrument_route("homework.grade_all", "homework")
def grade_all_submissions(homework_id: str):
    """
    Synchronous batch grading — grades every pending submission using Gemma 4.

    POST /api/homework/{homework_id}/grade-all

    Runs inline (no Cloud Run Job dispatch). Deploy with --timeout=540 to support
    classes up to ~40 students on Vertex AI.

    Flow:
      1. Verify submissions are closed (open_for_submission == False)
      2. Verify answer key has questions
      3. Fetch all pending submissions
      4. For each: Gemma 4 multimodal grade → annotate → upload → write Mark
      5. Send teacher push notification when complete
      6. Return { graded, errors, results }
    """
    teacher_id, err = require_role(request, "teacher")
    if err:
        return jsonify({"error": err}), 401

    homework = get_doc("answer_keys", homework_id)
    if not homework:
        return jsonify({"error": "Homework not found"}), 404
    if homework.get("teacher_id") != teacher_id:
        return jsonify({"error": "forbidden"}), 403

    if homework.get("open_for_submission", True):
        return jsonify({"error": "Close submissions before grading"}), 400

    if not homework.get("questions"):
        return jsonify({"error": "No answer key found. Add one before grading."}), 400

    class_id = homework.get("class_id", "")
    class_doc = get_doc("classes", class_id)
    education_level = homework.get("education_level") or (
        class_doc.get("education_level") if class_doc else "Form 4"
    )

    # ── Fetch pending submissions ─────────────────────────────────────────────
    pending = query("student_submissions", [
        ("answer_key_id", "==", homework_id),
        ("status", "==", "pending"),
    ])

    if not pending:
        return jsonify({"message": "No pending submissions to grade", "graded": 0}), 200

    # Mark homework as grading in progress
    upsert("answer_keys", homework_id, {
        "grading_status": "in_progress",
        "grading_started_at": _now_iso(),
    })

    # Defer heavy imports to avoid cold-start overhead on every request
    from shared.annotator import annotate_image
    from shared.gcs_client import generate_signed_url, upload_bytes
    from shared.gemma_client import grade_submission as _grade
    from shared.models import GradingVerdict, Mark

    graded = 0
    errors = 0
    results = []

    for sub in pending:
        sub_id = sub.get("id") or sub.get("submission_id")
        student_id = sub.get("student_id", "")
        image_urls: list = sub.get("image_urls") or []

        if not image_urls:
            logger.warning("grade-all: submission %s has no images — skipping", sub_id)
            upsert("student_submissions", sub_id, {
                "status": "error",
                "error": "no images attached",
            })
            errors += 1
            continue

        # Mark as "grading" so a polling teacher sees intermediate progress
        upsert("student_submissions", sub_id, {"status": "grading"})

        try:
            # Download primary image from GCS
            image_bytes = _download_image_from_url(image_urls[0])

            # Grade — Gemma 4 reads handwriting directly (single multimodal call)
            raw_verdicts = _grade(image_bytes, homework, education_level)
            verdicts = [GradingVerdict(**v) for v in raw_verdicts if isinstance(v, dict)]
            if not verdicts:
                logger.warning("[grade-all] Gemma returned 0 verdicts for sub=%s, skipping", sub_id)
                upsert("student_submissions", sub_id, {"status": "error", "error": "AI could not grade this submission"})
                results.append({"submission_id": sub_id, "status": "error", "error": "no verdicts"})
                continue
            score = sum(v.awarded_marks for v in verdicts)
            max_score = (
                sum(v.max_marks for v in verdicts)
                or float(homework.get("total_marks") or 1)
            )
            percentage = round(score / max_score * 100, 1) if max_score else 0.0

            # Annotate original image with ticks/crosses/scores
            verdicts_dicts = [v.model_dump() for v in verdicts]
            annotated_bytes = annotate_image(image_bytes, verdicts_dicts, None)

            # Upload annotated image to GCS marked bucket (private)
            blob_name = f"{student_id}/{uuid.uuid4()}.jpg"
            upload_bytes(settings.GCS_BUCKET_MARKED, blob_name, annotated_bytes, public=False)
            marked_url = generate_signed_url(settings.GCS_BUCKET_MARKED, blob_name, expiry_minutes=60 * 24 * 7)

            # Role invariants — refuse to save the Mark if either id
            # is wired to the wrong collection.
            from shared.role_invariants import assert_is_student, assert_is_teacher, RoleInvariantError
            try:
                assert_is_student(student_id)
                assert_is_teacher(teacher_id)
            except RoleInvariantError as e:
                logger.error("answer_keys close-and-grade role-invariant violated: %s", e)
                continue

            # Write Mark document (teacher review required before student can see)
            mark_doc = Mark(
                student_id=student_id,
                class_id=class_id,
                answer_key_id=homework_id,
                teacher_id=teacher_id,
                score=score,
                max_score=max_score,
                percentage=percentage,
                verdicts=verdicts,
                marked_image_url=marked_url,
                source="student_submission",
                approved=False,
            )
            upsert("marks", mark_doc.id, mark_doc.model_dump())

            # Update submission: status → graded, link mark
            now = _now_iso()
            upsert("student_submissions", sub_id, {
                "status": "graded",
                "mark_id": mark_doc.id,
                "score": score,
                "max_score": max_score,
                "percentage": percentage,
                "marked_image_url": marked_url,
                "graded_at": now,
                "grading_model": "gemma4-26b",
                "verdicts": verdicts_dicts,
            })

            graded += 1
            results.append({
                "submission_id": sub_id,
                "student_id": student_id,
                "score": score,
                "max_score": max_score,
                "percentage": percentage,
            })
            logger.info("grade-all: graded submission %s — %.1f%%", sub_id, percentage)

        except Exception:
            logger.exception("grade-all: failed to grade submission %s", sub_id)
            upsert("student_submissions", sub_id, {
                "status": "error",
                "error": "grading_failed",
            })
            errors += 1

    # ── Mark homework complete ────────────────────────────────────────────────
    upsert("answer_keys", homework_id, {
        "grading_status": "complete",
        "grading_completed_at": _now_iso(),
    })

    # ── Notify teacher ────────────────────────────────────────────────────────
    # ── Push notification ─────────────────────────────────────────────────────
    try:
        from functions.push import send_teacher_notification
        hw_title = homework.get("title", "Homework")
        body_msg = (
            f"{graded} submission(s) graded. {errors} could not be processed."
            if errors
            else f"All {graded} submission(s) graded successfully."
        )
        send_teacher_notification(
            teacher_id=teacher_id,
            title=f"Grading complete — {hw_title}",
            body=body_msg,
            data={
                "screen": "HomeworkDetail",
                "answer_key_id": homework_id,
                "class_id": homework.get("class_id", ""),
                "class_name": homework.get("title", hw_title),
            },
        )
    except Exception:
        logger.warning("grade-all: push notification failed (non-fatal)")

    # ── WhatsApp notification to teacher ──────────────────────────────────────
    # Sends a message to open the free 24-hour session window. The teacher can
    # then reply "results" to review and approve submissions via WhatsApp.
    try:
        teacher_doc = get_doc("teachers", teacher_id)
        if teacher_doc and teacher_doc.get("phone"):
            from shared.whatsapp_client import send_text as _wa_send
            hw_title = homework.get("title", "Homework")
            class_id = homework.get("class_id", "")
            class_doc = get_doc("classes", class_id)
            class_name = class_doc.get("name", "") if class_doc else ""
            label = f"{class_name} — {hw_title}" if class_name else hw_title
            if errors:
                summary = f"{graded} graded, {errors} error(s)."
            else:
                summary = f"All {graded} submission(s) graded."
            _wa_send(
                teacher_doc["phone"],
                f"Grading complete for *{label}* ✅\n"
                f"{summary}\n\n"
                "Reply *results* to review and approve submissions on WhatsApp.",
            )
    except Exception:
        logger.warning("grade-all: WhatsApp notification failed (non-fatal)")

    return jsonify({
        "graded": graded,
        "errors": errors,
        "results": results,
    }), 200


# ── POST /api/homework/{id}/approve-all ───────────────────────────────────────

@homework_bp.post("/homework/<homework_id>/approve-all")
@instrument_route("homework.approve_all", "homework")
def approve_all_submissions(homework_id: str):
    """
    Batch-approve all graded submissions for a homework.
    POST /api/homework/{homework_id}/approve-all
    """
    teacher_id, err = require_role(request, "teacher")
    if err:
        return jsonify({"error": err}), 401

    homework = get_doc("answer_keys", homework_id)
    if not homework:
        return jsonify({"error": "Homework not found"}), 404
    if homework.get("teacher_id") != teacher_id:
        return jsonify({"error": "forbidden"}), 403

    graded_subs = query("student_submissions", [
        ("answer_key_id", "==", homework_id),
        ("status", "==", "graded"),
    ])

    if not graded_subs:
        return jsonify({"message": "No graded submissions to approve", "approved": 0}), 200

    now = _now_iso()
    approved_count = 0
    for sub in graded_subs:
        sub_id = sub.get("id") or sub.get("submission_id")
        upsert("student_submissions", sub_id, {
            "status": "approved",
            "approved_at": now,
            "approved_by": teacher_id,
        })
        mark_id = sub.get("mark_id")
        if mark_id:
            upsert("marks", mark_id, {"approved": True, "approved_at": now})
        approved_count += 1
        # Update student weakness profile (fire and forget)
        approved_sub = {**sub, "status": "approved", "approved_at": now}
        update_student_weaknesses(sub.get("student_id", ""), approved_sub)

    return jsonify({
        "message": f"Approved {approved_count} submission(s)",
        "approved": approved_count,
    }), 200
